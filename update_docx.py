from docx import Document
import sys

def update_report(input_file, output_file):
    doc = Document(input_file)
    
    rep_map = {
        "Catálogo pesquisável com filtros por categoria e descrições geradas por IA.": "Catálogo pesquisável com filtros por categoria, com descrições geradas por IA e tradução automática multi-idioma (Google Translate API).",
        "Gestão de empréstimos com prazos, histórico e notificações in-app.": "Gestão de empréstimos com prazos, histórico e lembretes automáticos de 24 horas via email e notificações in-app.",
        "Autenticação por email/palavra-passe, magic link e recuperação de conta.": "Autenticação via Google, email institucional exclusivo (@agr-tc.pt), magic link e recuperação de conta.",
    }
    
    new_doc_point = "Documentação interativa com Termos e Condições e Política de Privacidade integrados, totalmente internacionalizados."
    inserted_new_point = False
    
    for para in doc.paragraphs:
        # Check for replacements
        for old_txt, new_txt in rep_map.items():
            if old_txt in para.text:
                # Update text directly. Note: This resets inline formatting (bold/italic) for this specific paragraph,
                # but keeps paragraph-level formatting (like bullets).
                para.text = para.text.replace(old_txt, new_txt)
        
        # Insert the new point after the authentication point
        if "Autenticação via Google" in para.text and not inserted_new_point:
            # We want to add a new paragraph with the same style
            new_p = para.insert_paragraph_before(new_doc_point, style=para.style)
            # Wait, insert_paragraph_before inserts BEFORE. Let's move the text around.
            # Actually we can just add the new point to the end of the authentication text if it's simpler, 
            # or use insert_paragraph_before on the NEXT paragraph.
            # But the simplest is just to add a new paragraph right before the next one.
            pass # We'll do this carefully below using index
            
    # To insert AFTER a paragraph, we need to find it and use insert_paragraph_before on the NEXT paragraph
    for i, para in enumerate(doc.paragraphs):
        if "Autenticação via Google" in para.text and not inserted_new_point:
            if i + 1 < len(doc.paragraphs):
                # Insert before the next paragraph
                new_p = doc.paragraphs[i+1].insert_paragraph_before(new_doc_point, style=para.style)
                inserted_new_point = True
                
    doc.save(output_file)
    print(f"Successfully updated and saved to {output_file}")

if __name__ == '__main__':
    update_report("Relatorio_do_Projeto.docx", "Relatorio_do_Projeto_Atualizado.docx")
