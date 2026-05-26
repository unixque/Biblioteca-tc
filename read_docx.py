from docx import Document
import sys

def dump_docx(filename, out_filename):
    try:
        doc = Document(filename)
        with open(out_filename, 'w', encoding='utf-8') as f:
            for para in doc.paragraphs:
                f.write(para.text + "\n")
            
            f.write("\n\n--- TABLES ---\n\n")
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    f.write(" | ".join(row_data) + "\n")
                f.write("\n---\n")
                
        print(f"Successfully dumped to {out_filename}")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    dump_docx("Relatorio_do_Projeto.docx", "docx_dump.txt")
